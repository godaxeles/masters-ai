"""Document processing service for extracting and chunking text."""

import hashlib
import uuid
from pathlib import Path
from typing import List, Optional, Tuple

import structlog
from pypdf import PdfReader
from docx import Document as DocxDocument

from config import settings
from models.document import DocumentChunk, ProcessedDocument
from utils.logging_config import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Process documents and extract text content."""

    def __init__(self):
        self.chunk_size = settings.max_chunk_size
        self.chunk_overlap = settings.chunk_overlap

    def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process a document and extract text chunks.

        Args:
            file_path: Path to the document file

        Returns:
            ProcessedDocument with extracted chunks

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        logger.info("Processing document", file_path=str(file_path))

        # Get file information
        file_stats = file_path.stat()
        file_type = file_path.suffix.lower().lstrip('.')

        # Extract text based on file type
        try:
            if file_type == 'pdf':
                text_content, total_pages = self._extract_pdf_text(file_path)
            elif file_type in ['txt', 'md']:
                text_content, total_pages = self._extract_text_file(file_path)
            elif file_type in ['docx', 'doc']:
                text_content, total_pages = self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Create chunks
            chunks = self._create_chunks(text_content, file_path.name)

            # Create processed document
            processed_doc = ProcessedDocument(
                file_path=file_path,
                file_name=file_path.name,
                file_size=file_stats.st_size,
                file_type=file_type,
                total_pages=total_pages,
                total_chunks=len(chunks),
                chunks=chunks
            )

            logger.info(
                "Document processed successfully",
                file_name=file_path.name,
                pages=total_pages,
                chunks=len(chunks),
                words=processed_doc.total_words
            )

            return processed_doc

        except Exception as e:
            logger.error(
                "Error processing document",
                file_path=str(file_path),
                error=str(e)
            )
            # Return document with error information
            return ProcessedDocument(
                file_path=file_path,
                file_name=file_path.name,
                file_size=file_stats.st_size,
                file_type=file_type,
                processing_errors=[str(e)]
            )

    def _extract_pdf_text(self, file_path: Path) -> Tuple[List[Tuple[str, int]], int]:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (list of (text, page_number) tuples, total_pages)
        """
        try:
            reader = PdfReader(str(file_path))
            text_pages = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():  # Only add non-empty pages
                        text_pages.append((text, page_num))
                except Exception as e:
                    logger.warning(
                        "Error extracting text from PDF page",
                        file_path=str(file_path),
                        page=page_num,
                        error=str(e)
                    )
                    continue

            total_pages = len(reader.pages)

            logger.info(
                "PDF text extracted",
                file_path=str(file_path),
                total_pages=total_pages,
                pages_with_text=len(text_pages)
            )

            return text_pages, total_pages

        except Exception as e:
            logger.error(
                "Error reading PDF file",
                file_path=str(file_path),
                error=str(e)
            )
            raise

    def _extract_text_file(self, file_path: Path) -> Tuple[List[Tuple[str, int]], int]:
        """Extract text from text file.

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (list of (text, page_number) tuples, total_pages)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # For text files, treat as single page
            return [(content, 0)], 1

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return [(content, 0)], 1
            except Exception as e:
                logger.error(
                    "Error reading text file",
                    file_path=str(file_path),
                    error=str(e)
                )
                raise

    def _extract_docx_text(self, file_path: Path) -> Tuple[List[Tuple[str, int]], int]:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (list of (text, page_number) tuples, total_pages)
        """
        try:
            doc = DocxDocument(str(file_path))

            # Extract all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            # Combine paragraphs
            content = '\n\n'.join(paragraphs)

            # Estimate pages (rough estimate: 500 words per page)
            word_count = len(content.split())
            estimated_pages = max(1, word_count // 500)

            return [(content, 0)], estimated_pages

        except Exception as e:
            logger.error(
                "Error reading DOCX file",
                file_path=str(file_path),
                error=str(e)
            )
            raise

    def _create_chunks(self, text_pages: List[Tuple[str, int]], source_file: str) -> List[DocumentChunk]:
        """Create text chunks from extracted content.

        Args:
            text_pages: List of (text, page_number) tuples
            source_file: Name of source file

        Returns:
            List of DocumentChunk objects
        """
        chunks = []

        for text, page_number in text_pages:
            if not text.strip():
                continue

            # Split text into sentences for better chunking
            sentences = self._split_into_sentences(text)

            # Create chunks with overlap
            page_chunks = self._create_page_chunks(sentences, page_number, source_file)
            chunks.extend(page_chunks)

        logger.info(
            "Text chunks created",
            source_file=source_file,
            total_chunks=len(chunks)
        )

        return chunks

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences.

        Args:
            text: Text to split

        Returns:
            List of sentences
        """
        # Simple sentence splitting (could be enhanced with NLTK)
        import re

        # Split on sentence endings, but keep common abbreviations together
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)

        # Clean up sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Filter very short sentences
                cleaned_sentences.append(sentence)

        return cleaned_sentences

    def _create_page_chunks(self, sentences: List[str], page_number: int, source_file: str) -> List[DocumentChunk]:
        """Create chunks from sentences with overlap.

        Args:
            sentences: List of sentences
            page_number: Page number for the chunks
            source_file: Source file name

        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        current_chunk = []
        current_word_count = 0

        for sentence in sentences:
            sentence_words = len(sentence.split())

            # If adding this sentence would exceed chunk size, create a chunk
            if current_word_count + sentence_words > self.chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunk = self._create_chunk(chunk_text, page_number, source_file)
                chunks.append(chunk)

                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(current_chunk)
                current_chunk = overlap_sentences + [sentence]
                current_word_count = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_word_count += sentence_words

        # Create final chunk if there's remaining content
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunk = self._create_chunk(chunk_text, page_number, source_file)
            chunks.append(chunk)

        return chunks

    def _get_overlap_sentences(self, sentences: List[str]) -> List[str]:
        """Get sentences for overlap with next chunk.

        Args:
            sentences: Current chunk sentences

        Returns:
            List of sentences for overlap
        """
        if not sentences:
            return []

        # Take last few sentences for overlap
        overlap_words = 0
        overlap_sentences = []

        for sentence in reversed(sentences):
            sentence_words = len(sentence.split())
            if overlap_words + sentence_words <= self.chunk_overlap:
                overlap_sentences.insert(0, sentence)
                overlap_words += sentence_words
            else:
                break

        return overlap_sentences

    def _create_chunk(self, text: str, page_number: int, source_file: str) -> DocumentChunk:
        """Create a DocumentChunk object.

        Args:
            text: Chunk text content
            page_number: Page number
            source_file: Source file name

        Returns:
            DocumentChunk object
        """
        # Generate unique chunk ID
        chunk_id = self._generate_chunk_id(text, source_file, page_number)

        # Calculate word and character counts
        word_count = len(text.split())
        char_count = len(text)

        return DocumentChunk(
            chunk_id=chunk_id,
            source_file=source_file,
            page_number=page_number,
            content=text,
            word_count=word_count,
            char_count=char_count,
            metadata={
                'source_file': source_file,
                'page': page_number,
                'chunk_id': chunk_id
            }
        )

    def _generate_chunk_id(self, text: str, source_file: str, page_number: int) -> str:
        """Generate a unique chunk identifier.

        Args:
            text: Chunk text
            source_file: Source file name
            page_number: Page number

        Returns:
            Unique chunk identifier
        """
        # Create deterministic ID based on content and metadata
        content_hash = hashlib.md5(text.encode()).hexdigest()[:8]
        return f"{source_file}_{page_number}_{content_hash}"

    def process_directory(self, directory_path: Path) -> List[ProcessedDocument]:
        """Process all documents in a directory.

        Args:
            directory_path: Path to directory containing documents

        Returns:
            List of ProcessedDocument objects
        """
        if not directory_path.exists() or not directory_path.is_dir():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        logger.info("Processing document directory", directory=str(directory_path))

        processed_docs = []
        supported_extensions = {'.pdf', '.txt', '.md', '.docx', '.doc'}

        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    processed_doc = self.process_document(file_path)
                    processed_docs.append(processed_doc)
                except Exception as e:
                    logger.error(
                        "Failed to process document",
                        file_path=str(file_path),
                        error=str(e)
                    )
                    continue

        logger.info(
            "Directory processing completed",
            directory=str(directory_path),
            total_documents=len(processed_docs),
            total_chunks=sum(doc.total_chunks for doc in processed_docs)
        )

        return processed_docs